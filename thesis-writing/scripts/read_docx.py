"""
read_docx.py —— 从 .docx 文件中提取内容与格式信息

用法：
  python read_docx.py <文件路径> [--mode content|format|all]
    content : 仅输出文本内容（带结构标记）
    format  : 仅输出格式检查报告
    all     : 同时输出内容和格式（默认）

依赖：pip install python-docx
"""

import argparse
import io
import sys
from pathlib import Path

# Windows 终端默认 GBK 编码，强制 stdout/stderr 使用 UTF-8 避免 UnicodeEncodeError
if sys.stdout.encoding and sys.stdout.encoding.lower() != "utf-8":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
if sys.stderr.encoding and sys.stderr.encoding.lower() != "utf-8":
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")

try:
    from docx import Document
    from docx.shared import Pt, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("错误：需要安装 python-docx 库。请运行：pip install python-docx", file=sys.stderr)
    sys.exit(1)


# ============================================================
# 工具函数
# ============================================================

def _pt(val):
    """将 EMU/Pt 值转为磅数字符串"""
    if val is None:
        return "未设置"
    return f"{val.pt:.1f}pt"


def _cm(val):
    """将 EMU 值转为厘米字符串"""
    if val is None:
        return "未设置"
    return f"{val.cm:.2f}cm"


def _font_desc(font):
    """描述一个 font 对象的主要属性"""
    parts = []
    if font.name:
        parts.append(font.name)
    if font.size:
        parts.append(_pt(font.size))
    if font.bold:
        parts.append("加粗")
    if font.italic:
        parts.append("斜体")
    return " | ".join(parts) if parts else "继承样式"


def _alignment_str(alignment):
    """段落对齐方式"""
    mapping = {
        WD_ALIGN_PARAGRAPH.LEFT: "左对齐",
        WD_ALIGN_PARAGRAPH.CENTER: "居中",
        WD_ALIGN_PARAGRAPH.RIGHT: "右对齐",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "两端对齐",
    }
    if alignment is None:
        return "未设置(继承)"
    return mapping.get(alignment, str(alignment))


# ============================================================
# 内容提取
# ============================================================

def extract_content(doc):
    """提取文档的结构化文本内容"""
    lines = []
    lines.append("=" * 60)
    lines.append("文档内容")
    lines.append("=" * 60)

    # 统计信息
    total_chars = 0
    heading_count = 0
    para_count = 0
    table_count = len(doc.tables)

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        text = para.text.strip()

        if not text:
            continue

        total_chars += len(text)

        if "Heading" in style_name or "标题" in style_name:
            # 尝试提取标题级别
            level = ""
            for ch in style_name:
                if ch.isdigit():
                    level = ch
                    break
            heading_count += 1
            prefix = f"[H{level}]" if level else "[H?]"
            lines.append(f"\n{prefix} {text}")
        else:
            para_count += 1
            lines.append(f"  {text}")

    lines.append("")
    lines.append("-" * 60)
    lines.append(f"统计：总字符数 ≈ {total_chars}，标题 {heading_count} 个，"
                 f"正文段落 {para_count} 个，表格 {table_count} 个")

    # 输出表格内容
    if table_count > 0:
        lines.append("")
        lines.append("=" * 60)
        lines.append("表格内容")
        lines.append("=" * 60)
        for i, table in enumerate(doc.tables, 1):
            lines.append(f"\n--- 表 {i} ({len(table.rows)}行 × {len(table.columns)}列) ---")
            for row_idx, row in enumerate(table.rows):
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                lines.append(f"  行{row_idx + 1}: {' | '.join(cells)}")

    return "\n".join(lines)


# ============================================================
# 格式检查
# ============================================================

def extract_format(doc):
    """提取文档的格式信息，用于合规检查"""
    lines = []
    lines.append("=" * 60)
    lines.append("格式检查报告")
    lines.append("=" * 60)

    # 1. 页面设置
    lines.append("\n## 页面设置")
    for i, section in enumerate(doc.sections):
        lines.append(f"  节 {i + 1}:")
        lines.append(f"    纸张大小: {_cm(section.page_width)} × {_cm(section.page_height)}")
        lines.append(f"    上边距: {_cm(section.top_margin)}")
        lines.append(f"    下边距: {_cm(section.bottom_margin)}")
        lines.append(f"    左边距: {_cm(section.left_margin)}")
        lines.append(f"    右边距: {_cm(section.right_margin)}")
        lines.append(f"    页眉距离: {_cm(section.header_distance)}")
        lines.append(f"    页脚距离: {_cm(section.footer_distance)}")

        # 页眉内容
        if section.header and not section.header.is_linked_to_previous:
            header_text = " ".join(p.text.strip() for p in section.header.paragraphs if p.text.strip())
            if header_text:
                lines.append(f"    页眉内容: {header_text}")
                # 页眉字体
                for p in section.header.paragraphs:
                    for run in p.runs:
                        if run.text.strip():
                            lines.append(f"    页眉字体: {_font_desc(run.font)}")
                            break
                    break

        # 页脚内容
        if section.footer and not section.footer.is_linked_to_previous:
            footer_text = " ".join(p.text.strip() for p in section.footer.paragraphs if p.text.strip())
            if footer_text:
                lines.append(f"    页脚内容: {footer_text}")

    # 2. 样式/格式统计
    lines.append("\n## 段落格式详情（按样式分组）")
    style_groups = {}
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else "(无样式)"
        if style_name not in style_groups:
            style_groups[style_name] = {
                "count": 0,
                "samples": [],
                "formats": []
            }
        group = style_groups[style_name]
        group["count"] += 1

        if len(group["samples"]) < 2:
            # 采样前40字符用于展示
            group["samples"].append(text[:40] + ("..." if len(text) > 40 else ""))

        if len(group["formats"]) < 2:
            fmt_info = {}
            # 段落级别
            pf = para.paragraph_format
            fmt_info["对齐"] = _alignment_str(pf.alignment)
            if pf.line_spacing:
                if pf.line_spacing_rule is not None:
                    fmt_info["行距"] = f"{pf.line_spacing} (规则: {pf.line_spacing_rule})"
                else:
                    fmt_info["行距"] = str(pf.line_spacing)
            if pf.space_before:
                fmt_info["段前"] = _pt(pf.space_before)
            if pf.space_after:
                fmt_info["段后"] = _pt(pf.space_after)
            if pf.first_line_indent:
                fmt_info["首行缩进"] = _pt(pf.first_line_indent)

            # Run 级别字体
            for run in para.runs:
                if run.text.strip():
                    fmt_info["字体"] = _font_desc(run.font)
                    break

            group["formats"].append(fmt_info)

    for style_name, group in sorted(style_groups.items()):
        lines.append(f"\n  [{style_name}] × {group['count']} 个段落")
        for sample in group["samples"]:
            lines.append(f"    示例: {sample}")
        for fmt in group["formats"]:
            for key, val in fmt.items():
                lines.append(f"    {key}: {val}")

    # 3. 图片数量
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_count += 1
    lines.append(f"\n## 其他信息")
    lines.append(f"  内嵌图片数量: {image_count}")
    lines.append(f"  表格数量: {len(doc.tables)}")

    return "\n".join(lines)


# ============================================================
# 主程序
# ============================================================

def main():
    parser = argparse.ArgumentParser(description="从 .docx 提取内容与格式信息")
    parser.add_argument("file", help="Word 文档路径（.docx）")
    parser.add_argument("--mode", choices=["content", "format", "all"],
                        default="all", help="提取模式（默认 all）")
    parser.add_argument("--output", "-o", help="输出到文件（UTF-8 编码），不指定则输出到终端")
    args = parser.parse_args()

    filepath = Path(args.file)
    if not filepath.exists():
        print(f"错误：文件不存在 - {filepath}", file=sys.stderr)
        sys.exit(1)
    if filepath.suffix.lower() != ".docx":
        print(f"警告：文件后缀不是 .docx，尝试按 docx 格式解析...", file=sys.stderr)

    doc = Document(str(filepath))

    parts = []
    if args.mode in ("content", "all"):
        parts.append(extract_content(doc))

    if args.mode == "all":
        parts.append("")

    if args.mode in ("format", "all"):
        parts.append(extract_format(doc))

    result = "\n".join(parts)

    if args.output:
        # 写入文件，避免终端编码问题
        out_path = Path(args.output)
        out_path.write_text(result, encoding="utf-8")
        print(f"已输出到: {out_path}")
    else:
        print(result)


if __name__ == "__main__":
    main()
